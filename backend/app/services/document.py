"""
Documentation Engine service.

Generates academic document sections using Claude (Anthropic API),
formats APA/IEEE citations, and exports to DOCX.
"""

import io
import logging
from datetime import datetime, timezone
from typing import Optional

import anthropic
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from app.config import get_settings
from app.models.document import AcademicDocument
from app.models.paper import Paper

logger = logging.getLogger(__name__)
settings = get_settings()

# ---------------------------------------------------------------------------
# Section definitions per academic field
# ---------------------------------------------------------------------------

SECTIONS_BY_FIELD: dict[str, list[str]] = {
    "computer_science": [
        "abstract",
        "introduction",
        "related_work",
        "methodology",
        "implementation",
        "evaluation",
        "conclusion",
        "references",
    ],
    "engineering": [
        "abstract",
        "introduction",
        "literature_review",
        "design_methodology",
        "implementation",
        "results_and_analysis",
        "discussion",
        "conclusion",
        "references",
    ],
    "business": [
        "executive_summary",
        "introduction",
        "literature_review",
        "research_methodology",
        "findings_and_analysis",
        "recommendations",
        "conclusion",
        "references",
    ],
    "health_sciences": [
        "abstract",
        "introduction",
        "literature_review",
        "methodology",
        "results",
        "discussion",
        "conclusion",
        "references",
    ],
}

SECTION_PROMPTS: dict[str, str] = {
    "abstract": (
        "Write a concise academic abstract (150-250 words) for this project. "
        "Summarise the problem, approach, and key findings."
    ),
    "introduction": (
        "Write a detailed introduction section (400-600 words). "
        "Explain the problem context, motivation, objectives, and document structure."
    ),
    "related_work": (
        "Write a related work / literature review section (500-800 words). "
        "Critically analyse prior work, identify gaps, and position this project."
    ),
    "literature_review": (
        "Write a comprehensive literature review (500-800 words). "
        "Survey relevant research, compare approaches, and identify research gaps."
    ),
    "methodology": (
        "Write a methodology section (400-700 words) describing the research design, "
        "data collection, analysis approach, and tools/technologies used."
    ),
    "design_methodology": (
        "Write a design and methodology section (500-800 words) covering "
        "system architecture, design decisions, and implementation strategy."
    ),
    "implementation": (
        "Write an implementation section (400-600 words) detailing the technical "
        "development process, key algorithms, and challenges overcome."
    ),
    "evaluation": (
        "Write an evaluation section (400-600 words) presenting experimental results, "
        "metrics, and performance analysis with relevant comparisons."
    ),
    "results_and_analysis": (
        "Write a results and analysis section (400-600 words) presenting findings, "
        "data interpretation, and discussion of outcomes."
    ),
    "results": (
        "Write a results section (400-600 words) presenting quantitative and qualitative "
        "findings with appropriate statistical analysis."
    ),
    "discussion": (
        "Write a discussion section (300-500 words) interpreting results, "
        "comparing with literature, and acknowledging limitations."
    ),
    "findings_and_analysis": (
        "Write a findings and analysis section (500-700 words) presenting key insights, "
        "data-driven conclusions, and business implications."
    ),
    "recommendations": (
        "Write a recommendations section (300-500 words) providing actionable "
        "suggestions based on the findings."
    ),
    "executive_summary": (
        "Write an executive summary (200-350 words) for business/management audiences. "
        "Cover the problem, approach, key findings, and recommendations."
    ),
    "research_methodology": (
        "Write a research methodology section (400-600 words) covering "
        "research design, data collection methods, and analytical framework."
    ),
    "conclusion": (
        "Write a conclusion section (300-450 words) summarising key contributions, "
        "lessons learned, and directions for future work."
    ),
    "references": (
        "List the references section header only — references are formatted separately."
    ),
}


# ---------------------------------------------------------------------------
# Claude section generation
# ---------------------------------------------------------------------------

def _build_system_prompt(project_title: str, field: str, papers: list[Paper]) -> str:
    paper_summaries = ""
    for i, p in enumerate(papers[:8], 1):
        authors_str = ", ".join(p.authors[:3]) + (" et al." if len(p.authors) > 3 else "")
        paper_summaries += (
            f"\n[{i}] {p.title} ({p.year or 'n.d.'}) — {authors_str}\n"
            f"    Abstract: {(p.abstract or '')[:200]}...\n"
        )

    return (
        f"You are an expert academic writer helping a student complete their final year project "
        f"titled '{project_title}' in the field of {field.replace('_', ' ').title()}.\n\n"
        f"Write in formal academic English. Be specific, cite the provided references where relevant "
        f"using [n] notation. Do NOT hallucinate references — only cite the numbered papers below.\n\n"
        f"Available references:\n{paper_summaries if paper_summaries else 'None provided — write without citations.'}"
    )


async def generate_section(
    document: AcademicDocument,
    section_name: str,
    project_title: str,
    field: str,
    papers: list[Paper],
    extra_context: Optional[str] = None,
) -> str:
    """Call Claude to generate a document section. Returns the generated text."""
    if not settings.ANTHROPIC_API_KEY:
        return (
            f"[{section_name.replace('_', ' ').title()} — AI generation requires ANTHROPIC_API_KEY]"
        )

    client = anthropic.Anthropic(api_key=settings.ANTHROPIC_API_KEY)
    system = _build_system_prompt(project_title, field, papers)
    section_instruction = SECTION_PROMPTS.get(
        section_name,
        f"Write the {section_name.replace('_', ' ')} section (400-600 words).",
    )
    user_msg = section_instruction
    if extra_context:
        user_msg += f"\n\nAdditional context from the student: {extra_context}"

    message = client.messages.create(
        model=settings.ANTHROPIC_MODEL,
        max_tokens=2048,
        system=system,
        messages=[{"role": "user", "content": user_msg}],
    )
    return message.content[0].text


# ---------------------------------------------------------------------------
# Citation formatting
# ---------------------------------------------------------------------------

def format_apa(paper: Paper) -> str:
    authors = paper.authors or ["Unknown"]
    if len(authors) == 1:
        author_str = authors[0]
    elif len(authors) == 2:
        author_str = f"{authors[0]} & {authors[1]}"
    else:
        author_str = f"{authors[0]} et al."
    year = paper.year or "n.d."
    doi_part = f" https://doi.org/{paper.doi}" if paper.doi else (f" {paper.url}" if paper.url else "")
    return f"{author_str} ({year}). {paper.title}.{doi_part}"


def format_ieee(paper: Paper, index: int) -> str:
    authors = paper.authors or ["Unknown"]
    author_str = ", ".join(authors[:3]) + (" et al." if len(authors) > 3 else "")
    year = paper.year or "n.d."
    doi_part = f", doi: {paper.doi}" if paper.doi else ""
    return f"[{index}] {author_str}, \"{paper.title},\" {year}{doi_part}."


def build_references_section(papers: list[Paper], citation_style: str) -> str:
    if not papers:
        return "No references indexed for this project."
    lines = []
    for i, paper in enumerate(papers, 1):
        if citation_style == "ieee":
            lines.append(format_ieee(paper, i))
        else:
            lines.append(format_apa(paper))
    return "\n\n".join(lines)


# ---------------------------------------------------------------------------
# DOCX export
# ---------------------------------------------------------------------------

def export_to_docx(document: AcademicDocument, papers: list[Paper]) -> bytes:
    """Generate a formatted Word document from the AcademicDocument sections."""
    doc = DocxDocument()

    # Title
    title_para = doc.add_heading(document.title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata line
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        f"Citation style: {document.citation_style.upper()} | "
        f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d')}"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()  # spacer

    sections = document.sections or {}
    for section_name, section_data in sections.items():
        if section_name == "references":
            continue  # handled below
        content = section_data.get("content", "") if isinstance(section_data, dict) else str(section_data)
        if not content:
            continue
        heading = section_name.replace("_", " ").title()
        doc.add_heading(heading, level=1)
        for para_text in content.split("\n\n"):
            para_text = para_text.strip()
            if para_text:
                doc.add_paragraph(para_text)

    # References section
    if papers:
        doc.add_heading("References", level=1)
        ref_text = build_references_section(papers, document.citation_style)
        for line in ref_text.split("\n\n"):
            if line.strip():
                doc.add_paragraph(line.strip())

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
